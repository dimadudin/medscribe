import json
import os
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

if getattr(sys, "frozen", False):
    if sys.stdout is None:
        sys.stdout = open(os.devnull, "w", encoding="utf-8")
    if sys.stderr is None:
        sys.stderr = open(os.devnull, "w", encoding="utf-8")

import numpy as np
import sounddevice as sd
from docx import Document
from docxtpl import DocxTemplate
from faster_whisper import WhisperModel
from huggingface_hub import snapshot_download
from tqdm import tqdm


def base_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).parent
    return Path(__file__).parent


TEMPLATES_DIR = base_dir() / "templates"
REPORTS_DIR = base_dir() / "reports"

MODEL_REPO = "Systran/faster-whisper-small"


def _progress_tqdm_class(on_percent: Callable[[int], None]) -> type[tqdm]:
    seen: dict[int, tuple[float, float]] = {}

    class _Progress(tqdm):
        def __init__(self, *args, **kwargs):
            if kwargs.get("file") is None and sys.stderr is None:
                kwargs["file"] = open(os.devnull, "w", encoding="utf-8")
            super().__init__(*args, **kwargs)

        def update(self, n=1):
            super().update(n)
            seen[id(self)] = (self.n or 0, self.total or 0)
            total = sum(t for _, t in seen.values())
            if total:
                done = sum(v for v, _ in seen.values())
                on_percent(min(100, int(done / total * 100)))

    return _Progress


class Recorder:
    def __init__(self, samplerate: int = 16000):
        self._samplerate = samplerate
        self._frames: list[np.ndarray] = []
        self._stream: sd.InputStream | None = None

    def start(self) -> None:
        self._frames = []

        def callback(indata, frame_count, time_info, status):
            self._frames.append(indata.copy())

        self._stream = sd.InputStream(
            samplerate=self._samplerate,
            channels=1,
            dtype="float32",
            callback=callback,
        )
        self._stream.start()

    def stop(self) -> np.ndarray:
        if self._stream is None:
            raise RuntimeError("Ошибка: запись не была запущена")
        self._stream.stop()
        self._stream.close()
        self._stream = None

        if not self._frames:
            raise RuntimeError("Ошибка: не удалось распознать аудио")

        return np.concatenate(self._frames, axis=0).flatten()


class Transcriber:
    def __init__(
        self,
        model_repo: str = MODEL_REPO,
        device: str = "auto",
        compute_type: str = "int8",
        log: Callable[[str], None] = print,
    ):
        self._model_repo = model_repo
        self._device = device
        self._compute_type = compute_type
        self._log = log
        self._model: WhisperModel | None = None

    def transcribe(self, audio: np.ndarray) -> str:
        if self._model is None:
            last_pct = -1

            def on_percent(pct: int) -> None:
                nonlocal last_pct
                if pct != last_pct:
                    last_pct = pct
                    self._log(f"Загрузка модели... {pct}%")

            self._log(f"Загрузка модели из Hugging Face ({self._model_repo})...")
            snapshot_download(
                self._model_repo, tqdm_class=_progress_tqdm_class(on_percent)
            )
            self._log("Модель загружена.")
            self._model = WhisperModel(
                self._model_repo,
                device=self._device,
                compute_type=self._compute_type,
            )
        self._log("Транскрипция...")
        segments, _ = self._model.transcribe(
            audio,
            language="ru",
            beam_size=5,
            vad_filter=True,
        )
        transcript = " ".join(segment.text for segment in segments).strip()
        self._log(f'Текст: "{transcript}"')
        return transcript


def discover_templates() -> dict[str, Path]:
    templates = _scan_templates(TEMPLATES_DIR) if TEMPLATES_DIR.is_dir() else {}
    if not templates:
        print(f"Шаблоны не найдены, создаю заглушку в {TEMPLATES_DIR}/{STUB_NAME}/...")
        create_stub_template()
        templates = _scan_templates(TEMPLATES_DIR)
    return templates


def _scan_templates(templates_dir: Path) -> dict[str, Path]:
    return {
        d.name: d
        for d in sorted(templates_dir.iterdir())
        if d.is_dir()
        and (d / "fields.toml").is_file()
        and (d / "template.docx").is_file()
    }


STUB_NAME = "рост-вес"

STUB_FIELDS_TOML = """\
[fields.HEIGHT]
aliases = ["рост"]
kind = "number"

[fields.WEIGHT]
aliases = ["вес"]
kind = "number"
"""

STUB_BSA_FORMULA = "{{ ((HEIGHT * WEIGHT / 3600) ** 0.5)|round(2) }}"


def create_stub_template(templates_dir: Path | None = None) -> Path:
    if templates_dir is None:
        templates_dir = TEMPLATES_DIR
    bundle = templates_dir / STUB_NAME
    bundle.mkdir(parents=True, exist_ok=True)
    (bundle / "fields.toml").write_text(STUB_FIELDS_TOML, encoding="utf-8")

    doc = Document()
    doc.add_heading("Антропометрия", level=1)
    doc.add_paragraph("Рост: {{ HEIGHT }} см")
    doc.add_paragraph("Вес: {{ WEIGHT }} кг")
    doc.add_paragraph(f"Площадь поверхности тела: {STUB_BSA_FORMULA} м²")
    doc.save(str(bundle / "template.docx"))
    return bundle


def save(
    findings: dict[str, float | str],
    bundle_dir: Path,
    reports_dir: Path = REPORTS_DIR,
    transcript: str | None = None,
) -> tuple[Path, Path]:
    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%d.%H_%M")
    tpl_path = bundle_dir / "template.docx"
    print(f"Загрузка шаблона из {tpl_path}...")
    reports_dir.mkdir(parents=True, exist_ok=True)

    payload: dict[str, float | str] = {}
    if transcript is not None:
        payload["transcript"] = transcript
    payload.update(findings)

    json_path = reports_dir / f"{timestamp}.json"
    print(f"Запись протокола в {json_path}...")
    json_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n")

    docx_path = reports_dir / f"{timestamp}.docx"
    print(f"Запись протокола в {docx_path}...")
    tpl = DocxTemplate(tpl_path)
    tpl.render(findings)
    tpl.save(str(docx_path))

    return json_path, docx_path
